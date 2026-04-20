import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"f:/Projects/AutoBill/output_docx"
os.makedirs(OUT_DIR, exist_ok=True)


def set_font(run, name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def add_p(doc, text="", align=None, bold=False, italic=False, size=11, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_label_value(doc, label, value, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(label)
    set_font(r1, size=size, bold=False)
    r2 = p.add_run(value)
    set_font(r2, size=size, bold=True)
    return p


def setup_page(doc, margin_cm=1.2):
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def set_fixed_layout(table, col_widths_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    total_twips = int(sum(col_widths_cm) * 567)
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w_cm in col_widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w_cm * 567)))
        tblGrid.append(gridCol)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcW")):
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(col_widths_cm[idx] * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def two_col_header(doc, left_lines, right_lines, left_bold_idx=(), right_bold_idx=(),
                   meta_lines=None):
    if meta_lines:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        col_widths = [5.5, 8.5, 4.5]
        set_fixed_layout(table, col_widths)
        cells = table.rows[0].cells
    else:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        col_widths = [8.5, 10.0]
        set_fixed_layout(table, col_widths)
        cells = table.rows[0].cells

    for c in cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)

    def fill(cell, lines, bold_idx, align, size=11):
        cell.paragraphs[0].text = ""
        for idx, line in enumerate(lines):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(line)
            set_font(r, size=size, bold=(idx in bold_idx))

    fill(cells[0], left_lines, left_bold_idx, WD_ALIGN_PARAGRAPH.CENTER)

    if meta_lines:
        fill(cells[1], right_lines, right_bold_idx, WD_ALIGN_PARAGRAPH.CENTER)
        fill(cells[2], meta_lines, (), WD_ALIGN_PARAGRAPH.LEFT, size=9)
    else:
        fill(cells[1], right_lines, right_bold_idx, WD_ALIGN_PARAGRAPH.CENTER)
    return table


def two_col_signature(doc, left_lines, right_lines):
    table = doc.add_table(rows=1, cols=2)
    cells = table.rows[0].cells
    for c in cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def fill(cell, lines):
        cell.paragraphs[0].text = ""
        for idx, (text, bold, italic, size) in enumerate(lines):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(text)
            set_font(r, size=size, bold=bold, italic=italic)

    fill(cells[0], left_lines)
    fill(cells[1], right_lines)
    return table


def top_info_row(doc, right_lines):
    table = doc.add_table(rows=1, cols=2)
    cells = table.rows[0].cells
    cells[0].width = Cm(12)
    cells[1].width = Cm(6)

    def fill_right(cell, lines):
        cell.paragraphs[0].text = ""
        for idx, line in enumerate(lines):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(line)
            set_font(r, size=10)

    fill_right(cells[1], right_lines)


def build_bvntw_page(doc, meta_lines):
    setup_page(doc)

    two_col_header(
        doc,
        [
            "BỘ Y TẾ",
            "BỆNH VIỆN NHI TRUNG ƯƠNG",
            "Số 18/879 Đường La Thành Quận",
            "Đống Đa - TP Hà Nội",
            "Khoa Ngoại Tiêu hóa - TT Ngoại Tổng hợp",
        ],
        [
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
            "Độc lập - Tự do - Hạnh phúc",
        ],
        left_bold_idx={1, 4},
        right_bold_idx={0, 1},
        meta_lines=meta_lines,
    )

    add_p(doc, "", space_after=6)
    add_p(doc, "GIẤY RA VIỆN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=6)

    info_table = doc.add_table(rows=1, cols=2)
    ic = info_table.rows[0].cells
    p1 = ic[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run("- Họ tên người bệnh: ")
    set_font(r, size=11)
    r = p1.add_run("LÊ KHÁNH CHI")
    set_font(r, size=11, bold=True)

    p2 = ic[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Ngày sinh: ")
    set_font(r, size=11)
    r = p2.add_run("07/06/2024")
    set_font(r, size=11, bold=True)
    r = p2.add_run("      Giới tính: ")
    set_font(r, size=11)
    r = p2.add_run("Nữ")
    set_font(r, size=11, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("- Dân tộc: ")
    set_font(r, size=11)
    r = p.add_run("Kinh")
    set_font(r, size=11, bold=True)
    r = p.add_run("      Nghề nghiệp: ")
    set_font(r, size=11)
    r = p.add_run("Sinh viên, học sinh")
    set_font(r, size=11, bold=True)

    add_label_value(doc, "- Mã số BHXH/Thẻ BHYT số: ", "TE1383825389141")
    add_label_value(doc, "- Địa chỉ: ", "Phường Đông Vệ, Thành phố Thanh Hóa, Thanh Hóa")
    add_label_value(doc, "- Vào viện lúc: ", "08 giờ 09 phút, ngày 03 tháng 10 năm 2024")
    add_label_value(doc, "- Ra viện lúc: ", "08 giờ 20 phút, ngày 15 tháng 10 năm 2024")
    add_label_value(doc, "- Chẩn đoán: ", "Hậu môn tiền đình/ giảm bạch cầu hạt/Tình trạng không có bạch cầu hạt")
    add_label_value(doc, "- Phương pháp điều trị: ", "Phẫu thuật")
    add_label_value(doc, "- Ghi chú: ", "Khám lại sau 2 tuần tại phòng khám số 23 lúc 14h chiều thứ 4, mua phiếu khám tại ô 1D.")
    add_label_value(doc, "+ Họ tên mẹ/bố/người giám hộ: ", "Lê Huyền Thanh/ Bố: Lê Khắc Tuấn")

    add_p(doc, "", space_after=4)
    add_p(doc, "Nếu có diễn biến bất thường đề nghị đến cơ sở Y tế gần nhất!",
          align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=8)

    two_col_signature(
        doc,
        [
            ("Ngày 15 Tháng 10 Năm 2024", False, False, 11),
            ("Thủ trưởng đơn vị", True, False, 11),
            ("(Ký tên, đóng dấu)", False, True, 10),
            ("TL. GIÁM ĐỐC", True, False, 11),
            ("PHÓ GIÁM ĐỐC TRUNG TÂM", True, False, 11),
            ("", False, False, 11),
            ("", False, False, 11),
            ("TS.BS. Trần Anh Quỳnh", True, False, 11),
        ],
        [
            ("Ngày 15 Tháng 10 Năm 2024", False, False, 11),
            ("Trưởng khoa", True, False, 11),
            ("", False, False, 11),
            ("", False, False, 11),
            ("", False, False, 11),
            ("", False, False, 11),
            ("", False, False, 11),
            ("TS.BS. Trần Anh Quỳnh", True, False, 11),
        ],
    )

    add_p(doc, "", space_after=4)
    add_p(doc, "15/10/2024  9:17:48AM", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)


def build_bvnth_page(doc):
    setup_page(doc)

    two_col_header(
        doc,
        [
            "SỞ Y TẾ THANH HÓA",
            "BỆNH VIỆN NHI THANH HÓA",
            "Khoa Răng hàm mặt",
        ],
        [
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
            "Độc lập - Tự do - Hạnh phúc",
        ],
        left_bold_idx={1, 2},
        right_bold_idx={0, 1},
        meta_lines=[
            "MS 02/TT25",
            "Mã BN: 24041628",
            "Số lưu trữ: 2680D0000464",
            "Mã YT: 38/38287/26/012456",
        ],
    )

    add_p(doc, "", space_after=6)
    add_p(doc, "GIẤY RA VIỆN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=6)

    info_table = doc.add_table(rows=1, cols=2)
    ic = info_table.rows[0].cells
    p1 = ic[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run("- Họ tên người bệnh: ")
    set_font(r, size=11)
    r = p1.add_run("LÊ KHÁNH CHI")
    set_font(r, size=11, bold=True)

    p2 = ic[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Nam/Nữ: ")
    set_font(r, size=11)
    r = p2.add_run("Nữ")
    set_font(r, size=11, bold=True)

    info2 = doc.add_table(rows=1, cols=2)
    ic2 = info2.rows[0].cells
    p1 = ic2[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run("- Ngày tháng năm sinh: ")
    set_font(r, size=11)
    r = p1.add_run("07/05/2024")
    set_font(r, size=11, bold=True)

    p2 = ic2[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Tuổi: ")
    set_font(r, size=11)
    r = p2.add_run("21 tháng")
    set_font(r, size=11, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("- Dân tộc: ")
    set_font(r, size=11)
    r = p.add_run("Kinh")
    set_font(r, size=11, bold=True)
    r = p.add_run("      Nghề nghiệp: ")
    set_font(r, size=11)
    r = p.add_run("Trẻ em dưới 6 tuổi")
    set_font(r, size=11, bold=True)

    add_label_value(doc, "- Số CCCD/CMND/Định danh công dân/Hộ chiếu: ", "038324011464")

    info3 = doc.add_table(rows=1, cols=2)
    ic3 = info3.rows[0].cells
    p1 = ic3[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run("- Mã số BHXH/Thẻ BHYT số: ")
    set_font(r, size=11)
    r = p1.add_run("TE 1 38 38 253 89141")
    set_font(r, size=11, bold=True)
    p2 = ic3[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Ngày cấp: ")
    set_font(r, size=11)

    add_label_value(doc, "- Địa chỉ: ", "57 Lạc Long Quân, Phường Hạc Thành, Tỉnh Thanh Hóa")
    add_label_value(doc, "- Vào viện lúc: ", "07 giờ 58 phút, Ngày 20 tháng 03 năm 2026")
    add_label_value(doc, "- Ra viện lúc: ", "15 giờ 30 phút, ngày 24 tháng 03 năm 2026")
    add_label_value(doc, "- Chẩn đoán: ", "K12-Viêm miệng và tổn thương liên quan, Viêm nướu răng")
    add_label_value(doc, "- Phương pháp điều trị: ", "Kháng sinh, hạ sốt, Chấm thuốc rửa miệng")
    add_label_value(doc, "- Ghi chú: ", "")
    add_label_value(doc, "- Người nhà: ", "Bố đẻ: LÊ KHẮC TUẤN, Mẹ đẻ: LÊ HUYỀN THANH, ĐT: 0912947608")

    add_p(doc, "", space_after=6)

    two_col_signature(
        doc,
        [
            ("Ngày 24 tháng 03 năm 2026", False, True, 11),
            ("Đại diện đơn vị", True, False, 11),
            ("(Ký, ghi rõ họ tên, đóng dấu)", False, True, 10),
            ("", False, False, 11),
            ("", False, False, 11),
            ("Người ký: Vũ Văn Thoan", True, False, 11),
        ],
        [
            ("Ngày 24 tháng 03 năm 2026", False, True, 11),
            ("Người hành nghề khám bệnh, chữa bệnh", True, False, 11),
            ("(Ký và ghi rõ họ tên)", False, True, 10),
            ("", False, False, 11),
            ("", False, False, 11),
            ("Người ký: Vũ Văn Thoan", True, False, 11),
        ],
    )


def make(filename, builder):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    builder(doc)
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print("Saved:", path)


page1_meta = [
    "MS: 01/BV-01",
    "Số bệnh án: 24.100617",
    "Mã y tế: 240241809",
    "In lần thứ: 1",
]

make("BENH VIEN NHI TRUNG UONG - Page 1.docx", lambda d: build_bvntw_page(d, page1_meta))
make("BENH VIEN NHI TRUNG UONG - Page 2.docx", lambda d: build_bvnth_page(d))
make("BENH VIEN NHI TRUNG UONG - Page 3.docx", lambda d: build_bvntw_page(d, page1_meta))
make("BENH VIEN NHI TRUNG UONG - Page 4.docx", lambda d: build_bvnth_page(d))
