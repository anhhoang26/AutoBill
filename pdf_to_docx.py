import os
import fitz
from docx import Document
from docx.shared import Inches, Emu
from docx.enum.section import WD_ORIENT

SRC_PDF = r"f:/Projects/AutoBill/BENH VIEN NHI TRUNG UONG.pdf"
OUT_DIR = r"f:/Projects/AutoBill/output_docx"
BASE_NAME = "BENH VIEN NHI TRUNG UONG"
DPI = 300

os.makedirs(OUT_DIR, exist_ok=True)

doc = fitz.open(SRC_PDF)
page_count = doc.page_count
print(f"[PDF2DOCX] Source has {page_count} pages, rendering at {DPI} DPI")

zoom = DPI / 72.0
matrix = fitz.Matrix(zoom, zoom)

for i in range(page_count):
    page = doc.load_page(i)
    rect = page.rect
    width_pt, height_pt = rect.width, rect.height
    width_in = width_pt / 72.0
    height_in = height_pt / 72.0

    pix = page.get_pixmap(matrix=matrix, alpha=False)
    img_path = os.path.join(OUT_DIR, f"__page_{i+1}.png")
    pix.save(img_path)

    docx_path = os.path.join(OUT_DIR, f"{BASE_NAME} - Page {i+1}.docx")
    docx = Document()
    section = docx.sections[0]
    section.page_width = Emu(int(width_in * 914400))
    section.page_height = Emu(int(height_in * 914400))
    section.top_margin = Emu(0)
    section.bottom_margin = Emu(0)
    section.left_margin = Emu(0)
    section.right_margin = Emu(0)
    section.header_distance = Emu(0)
    section.footer_distance = Emu(0)
    section.gutter = Emu(0)
    if width_pt > height_pt:
        section.orientation = WD_ORIENT.LANDSCAPE

    para = docx.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Emu(0)
    pf.space_after = Emu(0)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_in), height=Inches(height_in))

    docx.save(docx_path)
    print(f"[PDF2DOCX] Page {i+1} -> {docx_path} ({width_in:.2f}in x {height_in:.2f}in)")

    try:
        os.remove(img_path)
    except OSError:
        pass

doc.close()
print("[PDF2DOCX] Done.")
